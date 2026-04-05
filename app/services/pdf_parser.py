"""
PDF和文档解析模块
支持PDF和Word文档的解析
"""

import re
import logging
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from enum import Enum
from io import BytesIO

logger = logging.getLogger(__name__)


class PDFEngine(Enum):
    AUTO = "auto"
    PDFPLUMBER = "pdfplumber"
    PYMUPDF = "pymupdf"


@dataclass
class PDFParseResult:
    text: str
    pages: int
    metadata: Dict[str, Any] = field(default_factory=dict)
    tables: List[List[List[str]]] = field(default_factory=list)
    images_count: int = 0
    is_scanned: bool = False
    confidence: float = 1.0
    errors: List[str] = field(default_factory=list)


class PDFParser:
    def __init__(self, engine: PDFEngine = PDFEngine.AUTO):
        self.engine = engine
        self._pdfplumber_available = self._check_pdfplumber()
        self._pymupdf_available = self._check_pymupdf()
    
    def _check_pdfplumber(self) -> bool:
        try:
            import pdfplumber
            return True
        except ImportError:
            return False
    
    def _check_pymupdf(self) -> bool:
        try:
            import fitz
            return True
        except ImportError:
            return False
    
    def parse(self, file_content: bytes, password: str = None) -> PDFParseResult:
        if self.engine == PDFEngine.AUTO:
            if self._pdfplumber_available:
                return self._parse_with_pdfplumber(file_content, password)
            elif self._pymupdf_available:
                return self._parse_with_pymupdf(file_content, password)
            else:
                return PDFParseResult(
                    text="",
                    pages=0,
                    errors=["No PDF parsing library available. Install pdfplumber or pymupdf."]
                )
        elif self.engine == PDFEngine.PDFPLUMBER:
            return self._parse_with_pdfplumber(file_content, password)
        elif self.engine == PDFEngine.PYMUPDF:
            return self._parse_with_pymupdf(file_content, password)
        
        return PDFParseResult(text="", pages=0, errors=["Invalid engine specified"])
    
    def _parse_with_pdfplumber(self, file_content: bytes, password: str = None) -> PDFParseResult:
        try:
            import pdfplumber
        except ImportError:
            return PDFParseResult(
                text="",
                pages=0,
                errors=["pdfplumber not installed. Run: pip install pdfplumber"]
            )
        
        text_parts = []
        tables = []
        images_count = 0
        page_count = 0
        metadata = {}
        
        try:
            with pdfplumber.open(BytesIO(file_content), password=password) as pdf:
                page_count = len(pdf.pages)
                metadata = {
                    "author": pdf.metadata.get("Author", ""),
                    "title": pdf.metadata.get("Title", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                    "producer": pdf.metadata.get("Producer", ""),
                }
                
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
                    
                    page_tables = page.extract_tables()
                    if page_tables:
                        tables.extend(page_tables)
                    
                    images_count += len(page.images) if hasattr(page, 'images') else 0
            
            full_text = "\n\n".join(text_parts)
            
            is_scanned = self._detect_scanned_pdf(full_text, page_count)
            
            return PDFParseResult(
                text=full_text,
                pages=page_count,
                metadata=metadata,
                tables=tables,
                images_count=images_count,
                is_scanned=is_scanned,
                confidence=0.3 if is_scanned else 0.95
            )
            
        except Exception as e:
            logger.error(f"PDF parsing error with pdfplumber: {e}")
            return PDFParseResult(
                text="",
                pages=0,
                errors=[str(e)]
            )
    
    def _parse_with_pymupdf(self, file_content: bytes, password: str = None) -> PDFParseResult:
        try:
            import fitz
        except ImportError:
            return PDFParseResult(
                text="",
                pages=0,
                errors=["pymupdf not installed. Run: pip install pymupdf"]
            )
        
        text_parts = []
        tables = []
        images_count = 0
        page_count = 0
        metadata = {}
        
        try:
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            if password:
                doc.authenticate(password)
            
            page_count = len(doc)
            metadata = {
                "author": doc.metadata.get("author", ""),
                "title": doc.metadata.get("title", ""),
                "creator": doc.metadata.get("creator", ""),
                "producer": doc.metadata.get("producer", ""),
            }
            
            for page_num in range(page_count):
                page = doc[page_num]
                page_text = page.get_text()
                text_parts.append(page_text)
                
                images = page.get_images()
                images_count += len(images)
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            
            is_scanned = self._detect_scanned_pdf(full_text, page_count)
            
            return PDFParseResult(
                text=full_text,
                pages=page_count,
                metadata=metadata,
                tables=tables,
                images_count=images_count,
                is_scanned=is_scanned,
                confidence=0.3 if is_scanned else 0.95
            )
            
        except Exception as e:
            logger.error(f"PDF parsing error with pymupdf: {e}")
            return PDFParseResult(
                text="",
                pages=0,
                errors=[str(e)]
            )
    
    def _detect_scanned_pdf(self, text: str, page_count: int) -> bool:
        if not text or len(text.strip()) < 50:
            return True

        avg_chars_per_page = len(text) / max(page_count, 1)
        if avg_chars_per_page < 100:
            return True

        return False


class DocxParser:
    def __init__(self):
        self._python_docx_available = self._check_python_docx()
    
    def _check_python_docx(self) -> bool:
        try:
            from docx import Document
            return True
        except ImportError:
            return False
    
    def parse(self, file_content: bytes) -> Dict[str, Any]:
        if not self._python_docx_available:
            return {
                "text": "",
                "errors": ["python-docx not installed. Run: pip install python-docx"]
            }
        
        try:
            from docx import Document
            
            doc = Document(BytesIO(file_content))
            
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            full_text = "\n\n".join(text_parts)
            
            return {
                "text": full_text,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "errors": []
            }
            
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return {
                "text": "",
                "errors": [str(e)]
            }


class DocumentParser:
    def __init__(self, pdf_engine: PDFEngine = PDFEngine.AUTO):
        self.pdf_parser = PDFParser(engine=pdf_engine)
        self.docx_parser = DocxParser()
    
    def parse(self, file_content: bytes, filename: str = "", content_type: str = "") -> Dict[str, Any]:
        ext = self._get_extension(filename, content_type)
        
        if ext == "pdf":
            result = self.pdf_parser.parse(file_content)
            return {
                "text": result.text,
                "pages": result.pages,
                "metadata": result.metadata,
                "tables": result.tables,
                "images_count": result.images_count,
                "is_scanned": result.is_scanned,
                "confidence": result.confidence,
                "errors": result.errors,
                "format": "pdf"
            }
        
        elif ext in ["docx", "doc"]:
            result = self.docx_parser.parse(file_content)
            return {
                "text": result.get("text", ""),
                "paragraphs": result.get("paragraphs", 0),
                "tables": result.get("tables", 0),
                "errors": result.get("errors", []),
                "format": "docx"
            }
        
        elif ext == "txt":
            try:
                text = file_content.decode("utf-8")
            except UnicodeDecodeError:
                try:
                    text = file_content.decode("gbk")
                except UnicodeDecodeError:
                    text = file_content.decode("utf-8", errors="ignore")
            
            return {
                "text": text,
                "format": "txt",
                "errors": []
            }
        
        else:
            try:
                text = file_content.decode("utf-8", errors="ignore")
                if text.strip():
                    return {
                        "text": text,
                        "format": "unknown",
                        "errors": []
                    }
            except:
                pass
            
            return {
                "text": "",
                "format": ext,
                "errors": [f"Unsupported file format: {ext}"]
            }
    
    def _get_extension(self, filename: str, content_type: str) -> str:
        if filename:
            ext = filename.lower().split(".")[-1] if "." in filename else ""
            if ext in ["pdf", "docx", "doc", "txt"]:
                return ext
        
        content_type_map = {
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "application/msword": "doc",
            "text/plain": "txt",
        }
        
        return content_type_map.get(content_type, "")


document_parser = DocumentParser()
pdf_parser = PDFParser()
docx_parser = DocxParser()
