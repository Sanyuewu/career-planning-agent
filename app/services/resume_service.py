# -*- coding: utf-8 -*-
"""
简历解析服务
遵循v4规范：支持PDF/DOCX/图片格式
"""

import os
import base64
import logging
import tempfile
import time
from typing import Optional, List, Dict, Tuple, Union
from pathlib import Path
from io import BytesIO
from pydantic import BaseModel

logger = logging.getLogger(__name__)

from app.config import settings
from app.ai.llm_client import llm_client, render_prompt
from app.services.resume_enhanced import (
    resume_field_extractor,
    resume_quality_evaluator,
)
from app.services.pdf_parser import DocumentParser, PDFParseResult


class ResumeParseResult(BaseModel):
    """简历解析结果"""
    basic_info: Dict
    education: List[Dict]
    skills: List[str]
    internships: List[Dict]
    projects: List[Dict]
    certs: List[str]
    awards: List[str]
    career_intent: Optional[str] = None
    inferred_soft_skills: Dict
    completeness: float = 0.0
    missing_dims: List[str] = []
    parse_method: str = "llm"  # llm / rule_补全 / mixed


class ResumeService:
    """简历解析服务"""
    
    SUPPORTED_FORMATS = [".pdf", ".docx", ".jpg", ".jpeg", ".png", ".doc", ".txt"]
    MAX_FILE_SIZE_MB = 10
    
    def __init__(self):
        self.document_parser = DocumentParser()
    
    def validate_file(self, file_path: str) -> Tuple[bool, str]:
        """验证文件"""
        path = Path(file_path)
        
        if not path.exists():
            return False, "文件不存在"
        
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            return False, f"不支持的文件格式：{ext}，支持：{self.SUPPORTED_FORMATS}"
        
        size_mb = path.stat().st_size / 1024 / 1024
        if size_mb > self.MAX_FILE_SIZE_MB:
            return False, f"文件大小{size_mb:.1f}MB超过{self.MAX_FILE_SIZE_MB}MB限制"
        
        return True, "验证通过"
    
    def extract_text(self, file_path: str) -> str:
        """提取文本内容"""
        path = Path(file_path)
        ext = path.suffix.lower()
        logger.debug("extract_text: %s (ext=%s)", path.name, ext)

        if ext in [".jpg", ".jpeg", ".png"]:
            return self._extract_image(file_path)

        try:
            with open(file_path, "rb") as f:
                file_content = f.read()

            parse_result = self.document_parser.parse(file_content, path.name, "")

            if parse_result.get("errors"):
                logger.warning("文档解析警告 [%s]: %s", path.name, parse_result['errors'])

            text = parse_result.get("text", "")

            # PDF 兜底：文本为空，或乱码率过高（非中文+非ASCII字符占比>40%）时切换 fitz
            if ext == ".pdf":
                if not text.strip():
                    logger.info("PDF文字层为空，切换OCR: %s", path.name)
                    return self._ocr_pdf(file_path)
                total = len(text.replace(" ", "").replace("\n", ""))
                if total > 0:
                    import re as _re
                    readable = len(_re.findall(r'[\u4e00-\u9fff\u0020-\u007e]', text))
                    if readable / total < 0.6:
                        logger.info("PDF乱码率过高(%.1f%%)，切换fitz: %s", (1 - readable/total)*100, path.name)
                        fitz_text = self._extract_with_fitz(file_path)
                        if fitz_text and len(fitz_text.strip()) > len(text.strip()):
                            text = fitz_text

            return text
        except Exception as e:
            if ext == ".pdf":
                return self._ocr_pdf(file_path)
            raise ValueError(f"文件解析失败: {e}")
    
    def _extract_pdf(self, file_path: str) -> str:
        """提取PDF文本"""
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extract = page.extract_text()
                    if extract:
                        text += extract + "\n---\n"
            
            if not text.strip():
                return self._ocr_pdf(file_path)
            
            return text
        except Exception as e:
            return self._ocr_pdf(file_path)
    
    def _extract_with_fitz(self, file_path: str) -> str:
        """直接用 PyMuPDF 提取文本（对中文 PDF 编码更友好）"""
        try:
            import fitz
            doc = fitz.open(file_path)
            return "".join(page.get_text() for page in doc)
        except Exception:
            return ""

    def _ocr_image_bytes(self, img_bytes: bytes) -> str:
        """用 RapidOCR 对图片字节做 OCR，返回识别文字"""
        try:
            from rapidocr_onnxruntime import RapidOCR
            import numpy as np
            from PIL import Image
            import io

            img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
            img_array = np.array(img)
            ocr = RapidOCR()
            result, _ = ocr(img_array)
            if not result:
                return ""
            return "\n".join([line[1] for line in result if line and len(line) > 1])
        except Exception as e:
            raise ValueError(f"OCR识别失败: {e}")

    def _ocr_pdf(self, file_path: str) -> str:
        """PDF文字提取 → 扫描版PDF用RapidOCR逐页识别"""
        try:
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()

            if text.strip():
                return text

            # 文字层为空（扫描版PDF）→ 逐页渲染为图片做OCR
            ocr_parts = []
            for page in doc:
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat)
                img_bytes = pix.tobytes("jpeg")
                page_text = self._ocr_image_bytes(img_bytes)
                if page_text.strip():
                    ocr_parts.append(page_text)

            full_text = "\n\n".join(ocr_parts)
            if not full_text.strip():
                raise ValueError("OCR未能识别到任何文字，请确认PDF文件清晰可读")
            return full_text
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"PDF内容提取失败（文件可能加密或损坏）: {e}")
    
    def _extract_docx(self, file_path: str) -> str:
        """提取Word文本（段落 + 表格，覆盖表格排版简历）"""
        try:
            from docx import Document
            doc = Document(file_path)

            parts: list[str] = []

            # 1. 普通段落
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    parts.append(t)

            # 2. 表格中的每个单元格（很多简历用表格排版）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and t not in parts:
                            parts.append(t)

            return "\n".join(parts)
        except Exception as e:
            raise ValueError(f"Word解析失败: {e}")
    
    def _extract_image(self, file_path: str) -> str:
        """图片简历 → 用RapidOCR提取文字"""
        try:
            with open(file_path, "rb") as f:
                img_bytes = f.read()
            text = self._ocr_image_bytes(img_bytes)
            if not text.strip():
                raise ValueError("OCR未能识别到任何文字，请确认图片清晰可读")
            return text
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"图片读取失败（格式不支持或文件损坏）: {e}")
    
    async def parse_resume(self, file_path: str) -> ResumeParseResult:
        """解析简历"""
        logger.info("开始解析简历文件: %s", Path(file_path).name)
        t0 = time.monotonic()
        valid, msg = self.validate_file(file_path)
        if not valid:
            logger.warning("文件验证失败: %s", msg)
            raise ValueError(msg)

        text = self.extract_text(file_path)

        if not text.strip():
            logger.error("文本提取结果为空: %s", Path(file_path).name)
            raise ValueError("无法从文件中提取文本内容")

        result = await self._parse_text(text)
        logger.info("简历解析完成 [%.2fs]: 技能=%d 实习=%d 项目=%d 完整度=%.0f%%",
                    time.monotonic() - t0, len(result.skills),
                    len(result.internships), len(result.projects), result.completeness * 100)
        return result
    
    async def parse_resume_bytes(
        self,
        file_content: bytes,
        filename: str = "",
        content_type: str = ""
    ) -> ResumeParseResult:
        """解析简历二进制内容"""
        logger.info("开始解析简历字节流: filename=%s size=%dB", filename, len(file_content))
        t0 = time.monotonic()
        ext = Path(filename).suffix.lower() if filename else ""
        
        if ext in [".jpg", ".jpeg", ".png"]:
            text = self._ocr_image_bytes(file_content)
            if not text.strip():
                raise ValueError("OCR未能识别到任何文字，请确认图片清晰可读")
            return await self._parse_text(text)
        
        parse_result = self.document_parser.parse(file_content, filename, content_type)
        
        if parse_result.get("errors"):
            logger.warning("文档解析警告 [%s]: %s", filename, parse_result['errors'])

        text = parse_result.get("text", "")

        if not text.strip():
            logger.error("字节流文本提取为空: %s", filename)
            raise ValueError("无法从文件中提取文本内容")

        result = await self._parse_text(text)
        logger.info("字节流简历解析完成 [%.2fs]: %s", time.monotonic() - t0, filename)
        return result
    
    async def parse_text(self, text: str) -> ResumeParseResult:
        """解析简历文本"""
        return await self._parse_text(text)
    
    async def _parse_text(self, text: str) -> ResumeParseResult:
        """使用LLM解析文本 + 规则增强"""
        logger.debug("LLM解析简历文本: %d字符", len(text))
        t0 = time.monotonic()
        prompt = render_prompt("student_portrait_v1.jinja2", resume_text=text)

        try:
            result = await llm_client.chat_structured(
                prompt=prompt,
                output_model=ResumeParseResult,
                temperature=0.2,
                max_tokens=4096,
            )

            result_dict = result.model_dump()
            enhanced_dict = resume_quality_evaluator.enhance_parse_result(text, result_dict)
            
            for key, value in enhanced_dict.items():
                if hasattr(result, key):
                    setattr(result, key, value)

            result.completeness = self._calculate_completeness(result)
            result.missing_dims = self._find_missing_dims(result)
            result = self._post_validate_soft_skills(result)
            result = self._rule_fallback(result, text)

            logger.info("LLM解析耗时 %.2fs，开始质量校验", time.monotonic() - t0)
            quality_eval = resume_quality_evaluator.evaluate_parse_result(result.model_dump())
            if not quality_eval["is_valid"]:
                logger.warning("简历解析质量问题: %s", quality_eval['issues'])
                # 若完全没有提取到内容（姓名/教育/技能全空），视为解析失败
                if result.completeness == 0.0 and not result.basic_info.get("name") and not result.skills:
                    raise ValueError(
                        "AI未能从文档中识别有效简历内容，请确认文件清晰可读（建议使用文字版PDF或DOCX）"
                    )

            return result
        except Exception as e:
            err_msg = str(e)
            logger.error("简历文本解析失败 [%.2fs]: %s", time.monotonic() - t0, err_msg)
            if "JSON" in err_msg or "parse" in err_msg.lower():
                raise ValueError("AI解析简历内容失败（返回格式异常），请确认简历内容清晰完整后重试")
            elif "timeout" in err_msg.lower() or "连接" in err_msg:
                raise ValueError("AI服务连接超时，请稍后重试")
            elif "所有LLM供应商均不可用" in err_msg:
                raise ValueError("AI服务暂时不可用，请稍后重试")
            else:
                raise ValueError(f"简历解析失败: {err_msg}")

    async def _parse_image_vision(self, image_b64: str) -> ResumeParseResult:
        """使用DeepSeek视觉能力直接解析图片简历"""
        from app.ai.llm_client import _get_client_by_provider, LLMCallError

        prompt_text = (
            "请仔细阅读这张简历图片，提取其中所有信息，按照以下JSON格式输出，不要输出其他内容：\n"
            '{"basic_info":{"name":"","school":"","major":"","grade":"","phone":"","email":""},'
            '"education":[{"school":"","degree":"","major":"","start_year":"","end_year":"","gpa":null}],'
            '"skills":[],'
            '"internships":[{"company":"","role":"","start_date":"","end_date":"","duration_months":0,"description":""}],'
            '"projects":[{"name":"","description":"","tech_stack":[],"role":""}],'
            '"certs":[],'
            '"awards":[],'
            '"career_intent":"",'
            '"inferred_soft_skills":{},'
            '"completeness":0.0,'
            '"missing_dims":[]}'
        )

        try:
            client = _get_client_by_provider(settings.LLM_PROVIDER)
            response = await client.client.chat.completions.create(
                model=client.model,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt_text},
                        {"type": "image_url", "image_url": {
                            "url": f"data:image/jpeg;base64,{image_b64}"
                        }},
                    ],
                }],
                temperature=0.1,
                max_tokens=2000,
            )
            raw = response.choices[0].message.content or ""
            raw = raw.strip()
            if raw.startswith("```"):
                lines = raw.split("\n")
                raw = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

            result = ResumeParseResult.model_validate_json(raw.strip())
            result.completeness = self._calculate_completeness(result)
            result.missing_dims = self._find_missing_dims(result)
            result = self._post_validate_soft_skills(result)
            return result
        except Exception as e:
            raise ValueError(f"图片简历解析失败（DeepSeek视觉API）: {e}")

    def _calculate_completeness(self, result: ResumeParseResult) -> float:
        """计算完整度"""
        weights = {
            "basic_info": 0.15,
            "education": 0.15,
            "skills": 0.20,
            "internships": 0.20,
            "projects": 0.10,
            "certs": 0.05,
            "awards": 0.05,
            "inferred_soft_skills": 0.10,
        }
        
        score = 0.0
        
        if result.basic_info.get("name"):
            score += weights["basic_info"]
        if result.education:
            score += weights["education"]
        if result.skills:
            score += weights["skills"] * min(len(result.skills) / 5, 1)
        if result.internships:
            score += weights["internships"] * min(len(result.internships) / 2, 1)
        if result.projects:
            score += weights["projects"] * min(len(result.projects) / 2, 1)
        if result.certs:
            score += weights["certs"]
        if result.awards:
            score += weights["awards"]
        
        soft_skills = result.inferred_soft_skills or {}
        non_null_count = sum(1 for v in soft_skills.values() if v and v.get("score") is not None)
        score += weights["inferred_soft_skills"] * (non_null_count / 5)
        
        return round(score, 3)
    
    def _find_missing_dims(self, result: ResumeParseResult) -> List[str]:
        """找出缺失维度"""
        missing = []
        
        if not result.basic_info.get("name"):
            missing.append("姓名")
        if not result.education:
            missing.append("教育经历")
        if not result.skills:
            missing.append("专业技能")
        if not result.internships:
            missing.append("实习经历")
        if not result.projects:
            missing.append("项目经历")
        
        soft_skills = result.inferred_soft_skills or {}
        null_count = sum(1 for v in soft_skills.values() if not v or v.get("score") is None)
        if null_count > 2:
            missing.append("软技能评估")
        
        return missing

    def _post_validate_soft_skills(self, result: ResumeParseResult) -> ResumeParseResult:
        """软技能后校验：清除无实际证据支撑的推断分数，防止 LLM 凭空打分"""
        soft = result.inferred_soft_skills or {}

        has_team_evidence = bool(result.projects or result.internships)
        # 无团队/实习经历时，communication 和 teamwork 若无文字证据则清除
        for key in ("communication", "teamwork"):
            item = soft.get(key)
            if item and item.get("score") is not None and not has_team_evidence:
                if not item.get("evidence"):
                    soft[key] = {"score": None, "evidence": None}

        # 无获奖且无项目时，innovation 若无文字证据则清除
        has_innovation_evidence = bool(result.awards or result.projects)
        item = soft.get("innovation")
        if item and item.get("score") is not None and not has_innovation_evidence:
            if not item.get("evidence"):
                soft["innovation"] = {"score": None, "evidence": None}

        result.inferred_soft_skills = soft
        return result

    def _rule_fallback(self, result: ResumeParseResult, text: str) -> ResumeParseResult:
        """O2-b: LLM解析后的规则兜底补全——技能/学历/实习时长三项关键字段"""
        import re as _re
        rule_patched = False

        # 1. skills 为空 → 正则提取技术关键词
        if not result.skills:
            extracted = resume_field_extractor.extract_skills_enhanced(text)
            if extracted.value:
                result.skills = extracted.value
                rule_patched = True
                logger.info("O2-b规则补全：从简历文本提取到 %d 个技能", len(result.skills))

        # 2. education 缺 degree → 正则补全
        if result.education:
            for edu in result.education:
                if not edu.get("degree"):
                    degree_res = resume_field_extractor.extract_degree(text)
                    if degree_res.source != "default":
                        edu["degree"] = degree_res.value
                        rule_patched = True
                        logger.debug("O2-b规则补全：education.degree = %s", degree_res.value)

        # 3. internships 缺 duration_months → 从起止日期推算
        if result.internships:
            date_pattern = _re.compile(
                r'(\d{4}[./-]\d{1,2})\s*[-–—~至]*\s*(\d{4}[./-]\d{1,2}|至今|present)',
                _re.IGNORECASE
            )
            for intern in result.internships:
                if not intern.get("duration_months"):
                    start = intern.get("start_date", "")
                    end = intern.get("end_date", "")
                    if start:
                        months = resume_field_extractor._estimate_duration(start, end or "")
                        if months > 0:
                            intern["duration_months"] = months
                            rule_patched = True
                            logger.debug("O2-b规则补全：internship duration_months = %d", months)

        # 无论是否触发兜底，均对技能列表做去重+归一化
        if result.skills:
            result.skills = self._normalize_skills(result.skills)

        if rule_patched:
            result.parse_method = "mixed" if result.parse_method == "llm" else result.parse_method
            # 重新计算完整度
            result.completeness = self._calculate_completeness(result)
            result.missing_dims = self._find_missing_dims(result)

        return result

    @staticmethod
    def _normalize_skills(skills: list) -> list:
        """技能列表归一化：去重、移除括号说明、统一大小写处理"""
        import re as _re
        seen: dict = {}  # normalized_key -> original (优先保留第一次出现的)
        for s in skills:
            if not s or not isinstance(s, str):
                continue
            # 移除括号内的说明（"Python（精通）" → "Python"）
            clean = _re.sub(r'[（(][^)）]*[)）]', '', s).strip()
            # 统一特殊符号（C/C++ 保留，C++ 也保留）
            clean = clean.replace('（', '').replace('）', '').strip()
            if not clean:
                continue
            # 归一化键：全小写 + 去除空格，用于去重判断
            key = clean.lower().replace(' ', '').replace('-', '').replace('_', '')
            # 常见别名归一化：Python3→Python, Vue.js→Vue, Node.js→Node.js(保留)
            _alias_map = {
                'python3': 'Python', 'python2': 'Python',
                'vuejs': 'Vue', 'vue3': 'Vue3', 'vue2': 'Vue',
                'reactjs': 'React', 'reactnative': 'React Native',
                'nodejs': 'Node.js', 'javascript': 'JavaScript',
                'typescript': 'TypeScript', 'golang': 'Go',
                'postgresql': 'PostgreSQL', 'mysql8': 'MySQL', 'mysql5': 'MySQL',
            }
            if key in _alias_map:
                clean = _alias_map[key]
                key = clean.lower().replace(' ', '')
            if key not in seen:
                seen[key] = clean
        return list(seen.values())


resume_service = ResumeService()


def compute_parse_quality(result: "ResumeParseResponse") -> dict:
    """
    D-1: 简历解析质量评分
    返回 {score: 0-100, level: str, suggestions: [str]}
    """
    score = 0
    suggestions = []

    # 基础信息 20分
    bi = result.basic_info or {}
    if bi.get("name"):
        score += 5
    else:
        suggestions.append("未识别到姓名，建议在简历顶部清晰写明")
    if bi.get("email") or bi.get("phone"):
        score += 5
    else:
        suggestions.append("未识别到联系方式，建议添加邮箱或手机号")
    if result.education:
        score += 10
    else:
        suggestions.append('未识别到教育经历，建议使用"教育背景"标题明确标注')

    # 技能 20分
    skill_count = len(result.skills or [])
    if skill_count >= 5:
        score += 20
    elif skill_count >= 3:
        score += 12
        suggestions.append(f"技能项较少（{skill_count}项），建议补充更多技术栈")
    elif skill_count > 0:
        score += 5
        suggestions.append("技能项不足（<3项），建议详细列举掌握的技术工具")
    else:
        suggestions.append('未识别到技能，建议单独添加"技能特长"一节')

    # 实习经历 25分
    internship_count = len(result.internships or [])
    if internship_count >= 2:
        score += 25
    elif internship_count == 1:
        score += 15
        suggestions.append("只有1段实习，建议补充更多实践经历或项目经历")
    else:
        suggestions.append("未识别到实习经历，实习经历对求职匹配度影响较大")

    # 项目经历 20分
    project_count = len(result.projects or [])
    if project_count >= 2:
        score += 20
    elif project_count == 1:
        score += 12
    else:
        suggestions.append("项目经历不足，建议补充课程项目或个人作品")

    # 求职意向 10分
    if result.career_intent:
        score += 10
    else:
        suggestions.append("未识别到求职意向，建议在简历头部写明目标岗位")

    # 证书/奖项 5分
    if result.certs or result.awards:
        score += 5

    score = min(score, 100)
    level = "优秀" if score >= 85 else "良好" if score >= 70 else "一般" if score >= 50 else "待完善"
    return {"score": score, "level": level, "suggestions": suggestions}


class ResumeParseResponse(BaseModel):
    """简历解析完整响应：ParseResult + Portrait + 换岗机会"""
    student_id: str
    basic_info: Dict
    education: List[Dict]
    skills: List[str]
    internships: List[Dict]
    projects: List[Dict]
    certs: List[str]
    awards: List[str]
    career_intent: Optional[str] = None
    inferred_soft_skills: Dict
    completeness: float
    missing_dims: List[str]
    competitiveness: float
    competitiveness_level: str
    transfer_opportunities: List[Dict] = []
    gap_mapped_transfers: List[Dict] = []
    parse_quality: Optional[Dict] = None   # D-1: 解析质量评分


class ResumeUploadResponse(BaseModel):
    task_id: str
    message: str
    result: Optional[ResumeParseResponse] = None
