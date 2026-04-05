# -*- coding: utf-8 -*-
"""
报告导出服务
支持PDF(HTML转换)和Word导出
"""

import io
import os
from datetime import datetime
from typing import Optional, List, Dict
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.config import DATA_DIR

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False


class ReportExportService:
    """报告导出服务"""
    
    def __init__(self):
        self.upload_dir = DATA_DIR / "reports"
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self._init_pdf_fonts()
    
    def _init_pdf_fonts(self):
        """初始化PDF中文字体"""
        if not HAS_REPORTLAB:
            return
        
        font_paths = [
            "C:/Windows/Fonts/msyh.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/simsun.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ]
        
        self.chinese_font = None
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    font_name = "ChineseFont"
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    self.chinese_font = font_name
                    break
                except Exception:
                    continue
    
    def export_to_pdf(
        self,
        report_data: Dict,
        output_path: Optional[str] = None
    ) -> str:
        """导出为PDF文档"""
        if not HAS_REPORTLAB:
            return self.export_to_html(report_data)
        
        if output_path is None:
            output_path = str(self.upload_dir / f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        if self.chinese_font:
            title_style = ParagraphStyle(
                'ChineseTitle',
                parent=styles['Title'],
                fontName=self.chinese_font,
                fontSize=24,
                spaceAfter=30,
            )
            heading_style = ParagraphStyle(
                'ChineseHeading',
                parent=styles['Heading1'],
                fontName=self.chinese_font,
                fontSize=16,
                spaceAfter=12,
                spaceBefore=20,
            )
            normal_style = ParagraphStyle(
                'ChineseNormal',
                parent=styles['Normal'],
                fontName=self.chinese_font,
                fontSize=11,
                leading=18,
            )
        else:
            title_style = styles['Title']
            heading_style = styles['Heading1']
            normal_style = styles['Normal']
        
        story = []
        
        story.append(Paragraph("职业规划报告", title_style))
        story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", normal_style))
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("基本信息", heading_style))
        
        basic_data = [
            ['项目', '内容'],
            ['姓名', report_data.get('student_name', '')],
            ['目标岗位', report_data.get('job_name', '')],
            ['综合匹配度', f"{report_data.get('overall_score', 0):.1f}%"],
        ]
        
        basic_table = Table(basic_data, colWidths=[4*cm, 10*cm])
        basic_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.4, 0.5, 0.9)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), self.chinese_font or 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(basic_table)
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("匹配维度分析", heading_style))
        
        dimensions = report_data.get("dimensions", {})
        dim_data = [
            ['维度', '得分', '说明'],
            ['基础要求', f"{dimensions.get('basic_requirements', {}).get('score', 0):.1f}%", 
             dimensions.get('basic_requirements', {}).get('detail', '')[:30]],
            ['职业技能', f"{dimensions.get('professional_skills', {}).get('score', 0):.1f}%",
             f"匹配{len(dimensions.get('professional_skills', {}).get('matched_skills', []))}项技能"],
            ['职业素养', f"{dimensions.get('professional_qualities', {}).get('score', 0):.1f}%",
             dimensions.get('professional_qualities', {}).get('detail', '')[:30]],
            ['发展潜力', f"{dimensions.get('development_potential', {}).get('score', 0):.1f}%",
             dimensions.get('development_potential', {}).get('detail', '')[:30]],
        ]
        
        dim_table = Table(dim_data, colWidths=[3*cm, 2*cm, 9*cm])
        dim_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.4, 0.5, 0.9)),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), self.chinese_font or 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        story.append(dim_table)
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("技能差距分析", heading_style))
        skill_gaps = report_data.get("skill_gaps", [])
        
        if skill_gaps:
            for gap in skill_gaps:
                skill_text = f"• {gap.get('skill', '')} - {gap.get('suggestion', '')}"
                story.append(Paragraph(skill_text, normal_style))
        else:
            story.append(Paragraph("暂无明显技能差距", normal_style))
        
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("行动计划", heading_style))
        action_plan = report_data.get("action_plan", [])
        
        if action_plan:
            for i, phase in enumerate(action_plan):
                phase_title = f"阶段{i+1}: {phase.get('phase', '')} ({phase.get('timeline', '')})"
                story.append(Paragraph(phase_title, normal_style))
                
                goals = phase.get("goals", [])
                for goal in goals:
                    story.append(Paragraph(f"  • {goal}", normal_style))
                story.append(Spacer(1, 10))
        else:
            story.append(Paragraph("暂无行动计划", normal_style))
        
        story.append(Spacer(1, 20))
        
        story.append(Paragraph("职业发展路径", heading_style))
        career_path = report_data.get("career_path", [])
        
        if career_path:
            for i, path in enumerate(career_path):
                path_text = f"{i+1}. {path.get('title', path.get('job', ''))}"
                if path.get('description'):
                    path_text += f" - {path.get('description')}"
                story.append(Paragraph(path_text, normal_style))
        else:
            story.append(Paragraph("暂无职业路径规划", normal_style))
        
        story.append(Spacer(1, 30))
        story.append(Paragraph("— 报告由AI职业规划智能体生成 —", normal_style))
        
        doc.build(story)
        return output_path
    
    def export_to_word(
        self,
        report_data: Dict,
        output_path: Optional[str] = None
    ) -> str:
        """导出为Word文档"""
        doc = Document()
        
        title = doc.add_heading("职业规划报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("")
        
        doc.add_heading("基本信息", level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        cells = table.rows[0].cells
        cells[0].text = "姓名"
        cells[1].text = report_data.get("student_name", "")
        
        cells = table.rows[1].cells
        cells[0].text = "目标岗位"
        cells[1].text = report_data.get("job_name", "")
        
        cells = table.rows[2].cells
        cells[0].text = "综合匹配度"
        cells[1].text = f"{report_data.get('overall_score', 0):.1f}%"
        
        doc.add_paragraph("")
        
        doc.add_heading("匹配维度分析", level=1)
        
        dimensions = report_data.get("dimensions", {})
        dim_table = doc.add_table(rows=5, cols=3)
        dim_table.style = 'Table Grid'
        
        headers = dim_table.rows[0].cells
        headers[0].text = "维度"
        headers[1].text = "得分"
        headers[2].text = "说明"
        
        dim_data = [
            ("基础要求", dimensions.get("basic_requirements", {}).get("score", 0), 
             dimensions.get("basic_requirements", {}).get("detail", "")),
            ("职业技能", dimensions.get("professional_skills", {}).get("score", 0),
             f"匹配{len(dimensions.get('professional_skills', {}).get('matched_skills', []))}项技能"),
            ("职业素养", dimensions.get("professional_qualities", {}).get("score", 0),
             dimensions.get("professional_qualities", {}).get("detail", "")),
            ("发展潜力", dimensions.get("development_potential", {}).get("score", 0),
             dimensions.get("development_potential", {}).get("detail", "")),
        ]
        
        for i, (name, score, detail) in enumerate(dim_data):
            row = dim_table.rows[i + 1].cells
            row[0].text = name
            row[1].text = f"{score:.1f}%"
            row[2].text = detail[:50] + "..." if len(detail) > 50 else detail

        # E-1: Word 中追加文本形式的雷达图（替代不可嵌入的ECharts）
        dim_score_dict = {name: score for name, score, _ in dim_data}
        doc.add_paragraph(self._dim_table_text(dim_score_dict))

        doc.add_paragraph("")
        
        doc.add_heading("技能差距分析", level=1)
        skill_gaps = report_data.get("skill_gaps", [])
        
        if skill_gaps:
            for gap in skill_gaps:
                p = doc.add_paragraph()
                p.add_run(f"• {gap.get('skill', '')}").bold = True
                p.add_run(f" - {gap.get('suggestion', '')}")
                if gap.get("jd_source"):
                    p.add_run(f" (来源: {gap.get('jd_source')})").italic = True
        else:
            doc.add_paragraph("暂无明显技能差距")
        
        doc.add_paragraph("")
        
        doc.add_heading("行动计划", level=1)
        action_plan = report_data.get("action_plan", [])
        
        if action_plan:
            for i, phase in enumerate(action_plan):
                doc.add_heading(f"阶段{i+1}: {phase.get('phase', '')}", level=2)
                doc.add_paragraph(f"时间线: {phase.get('timeline', '')}")
                
                goals = phase.get("goals", [])
                for goal in goals:
                    doc.add_paragraph(f"• {goal}", style='List Bullet')
        else:
            doc.add_paragraph("暂无行动计划")
        
        doc.add_paragraph("")

        doc.add_heading("职业发展路径", level=1)
        career_path = report_data.get("career_path", [])

        if career_path:
            for path in career_path:
                p = doc.add_paragraph()
                p.add_run(f"• {path.get('title', path.get('job', ''))}").bold = True
                p.add_run(f" - {path.get('description', '')}")
        else:
            doc.add_paragraph("暂无职业路径规划")

        # 能力维度文字表格（替代雷达图）
        dimensions = report_data.get("dimensions", {})
        doc.add_paragraph("")
        doc.add_heading("能力维度得分（雷达图数据）", level=1)
        dim_score_table = doc.add_table(rows=5, cols=2)
        dim_score_table.style = 'Table Grid'
        dim_rows = [
            ("基础要求", f"{dimensions.get('basic_requirements', {}).get('score', 0):.1f}%"),
            ("职业技能", f"{dimensions.get('professional_skills', {}).get('score', 0):.1f}%"),
            ("职业素养", f"{dimensions.get('professional_qualities', {}).get('score', 0):.1f}%"),
            ("发展潜力", f"{dimensions.get('development_potential', {}).get('score', 0):.1f}%"),
            ("市场需求", f"{dimensions.get('market_demand', {}).get('score', 0) if dimensions.get('market_demand') else 0:.1f}%"),
        ]
        for i, (dim_name, dim_score) in enumerate(dim_rows):
            row = dim_score_table.rows[i].cells
            row[0].text = dim_name
            row[1].text = dim_score

        # 6章节详细内容
        chapters = report_data.get("chapters", [])
        if chapters:
            doc.add_paragraph("")
            doc.add_heading("报告详细内容", level=1)
            for ch in chapters:
                ch_title = ch.get("title", "")
                ch_icon = ch.get("icon", "")
                ch_content = ch.get("content_md", "") or ch.get("content", "")
                action_items = ch.get("action_items", [])
                doc.add_heading(f"{ch_icon} {ch_title}", level=2)
                if ch_content:
                    import re
                    # 去掉 Markdown 格式符号，保留纯文本
                    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', ch_content)
                    clean = re.sub(r'#+\s*', '', clean)
                    for line in clean.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                if action_items:
                    for item in action_items:
                        t = item.get("title", "") or item.get("phase", "")
                        tl = item.get("timeline", "")
                        desc = item.get("description", "")
                        ver = item.get("verification", "")
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{t}").bold = True
                        if tl:
                            p.add_run(f" （{tl}）")
                        if desc:
                            doc.add_paragraph(f"  {desc}")
                        if ver:
                            doc.add_paragraph(f"  ✓ 验收：{ver}")

        doc.add_paragraph("")
        doc.add_paragraph("以下内容由AI生成，仅供参考，关键决策建议结合专业顾问意见。").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("— 报告由AI职业规划智能体生成 —").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if output_path is None:
            output_path = str(self.upload_dir / f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        
        doc.save(output_path)
        return output_path
    
    def _build_chapters_html(self, chapters: List[Dict]) -> str:
        """将6章节内容渲染为HTML段落"""
        if not chapters:
            return ""
        html = ""
        for ch in chapters:
            title = ch.get("title", "")
            content = ch.get("content_md", "") or ch.get("content", "")
            action_items = ch.get("action_items", [])
            icon = ch.get("icon", "")
            html += f'<h2>{icon} {title}</h2>\n'
            if content:
                # 简单 Markdown → HTML：换行、**粗体**
                import re
                c = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                c = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', c)
                c = re.sub(r'\n#+\s*(.+)', r'<br/><strong>\1</strong>', c)
                c = c.replace("\n", "<br/>")
                html += f'<p>{c}</p>\n'
            if action_items:
                html += '<ul class="action-items">\n'
                for item in action_items:
                    t = item.get("title", "") or item.get("phase", "")
                    tl = item.get("timeline", "")
                    desc = item.get("description", "")
                    ver = item.get("verification", "")
                    html += f'<li><strong>{t}</strong>'
                    if tl:
                        html += f' <span class="timeline">（{tl}）</span>'
                    if desc:
                        html += f'<br/><span class="item-desc">{desc}</span>'
                    if ver:
                        html += f'<br/><span class="item-ver">✓ 验收：{ver}</span>'
                    html += '</li>\n'
                html += '</ul>\n'
        return html

    def export_to_html(
        self,
        report_data: Dict
    ) -> str:
        """导出为HTML（可用于打印PDF），内嵌ECharts CDN保持图表可交互"""
        dimensions = report_data.get("dimensions", {})
        skill_gaps = report_data.get("skill_gaps", [])
        action_plan = report_data.get("action_plan", [])
        career_path = report_data.get("career_path", [])
        chapters = report_data.get("chapters", [])

        skill_gaps_html = ""
        for gap in skill_gaps:
            skill_gaps_html += f"""
            <div class="gap-item">
                <strong>{gap.get('skill', '')}</strong>
                <p>{gap.get('suggestion', '')}</p>
                <small>来源: {gap.get('jd_source', '')}</small>
            </div>
            """
        if not skill_gaps_html:
            skill_gaps_html = "<p>暂无明显技能差距</p>"

        action_plan_html = ""
        for i, phase in enumerate(action_plan):
            goals_html = "".join([f"<li>{g}</li>" for g in phase.get("goals", [])])
            desc = phase.get("description", "")
            title = phase.get("title", "") or phase.get("phase", "")
            tl = phase.get("timeline", "")
            action_plan_html += f"""
            <div class="phase">
                <h3>阶段{i+1}: {title} {f'（{tl}）' if tl else ''}</h3>
                {f'<p>{desc}</p>' if desc else ''}
                {f'<ul>{goals_html}</ul>' if goals_html else ''}
            </div>
            """
        if not action_plan_html:
            action_plan_html = "<p>暂无行动计划</p>"

        career_path_html = ""
        for path in career_path:
            career_path_html += f"""
            <p><strong>{path.get('title', path.get('job', ''))}</strong> - {path.get('description', '')}</p>
            """
        if not career_path_html:
            career_path_html = "<p>暂无职业路径规划</p>"

        # 雷达图数据（用于ECharts）
        dim_data = {
            "基础要求": dimensions.get("basic_requirements", {}).get("score", 0),
            "职业技能": dimensions.get("professional_skills", {}).get("score", 0),
            "职业素养": dimensions.get("professional_qualities", {}).get("score", 0),
            "发展潜力": dimensions.get("development_potential", {}).get("score", 0),
            "市场需求": dimensions.get("market_demand", {}).get("score", 0) if dimensions.get("market_demand") else 0,
        }
        import json as _json
        radar_indicator = _json.dumps([{"name": k, "max": 100} for k in dim_data.keys()])
        radar_values = _json.dumps(list(dim_data.values()))

        # E-1: 服务端SVG雷达图（不依赖CDN，离线/PDF场景可用）
        radar_svg = self._build_radar_svg(dim_data)

        # 6章节内容
        chapters_html = self._build_chapters_html(chapters) if chapters else ""

        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>职业规划报告</title>
    <script src="https://cdn.jsdelivr.net/npm/echarts@5/dist/echarts.min.js" onerror="document.getElementById('radarChart').style.display='none';document.getElementById('radarSvg').style.display='block'"></script>
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 900px; margin: 0 auto; padding: 24px; color: #333; }}
        h1 {{ text-align: center; color: #1890ff; margin-bottom: 4px; }}
        h2 {{ color: #333; border-bottom: 2px solid #1890ff; padding-bottom: 5px; margin-top: 32px; }}
        h3 {{ color: #555; }}
        table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background-color: #f5f5f5; font-weight: 600; }}
        .score {{ font-size: 28px; color: #1890ff; font-weight: bold; }}
        .tag {{ display: inline-block; background: #e6f7ff; padding: 2px 8px; margin: 2px; border-radius: 4px; font-size: 13px; }}
        .gap-item {{ margin: 10px 0; padding: 10px 14px; background: #fff7e6; border-left: 3px solid #faad14; border-radius: 4px; }}
        .phase {{ margin: 15px 0; padding: 14px 16px; background: #f6ffed; border-radius: 8px; border-left: 3px solid #52c41a; }}
        .action-items {{ padding-left: 20px; }}
        .action-items li {{ margin: 8px 0; }}
        .timeline {{ color: #888; font-size: 13px; }}
        .item-desc {{ color: #555; font-size: 13px; }}
        .item-ver {{ color: #52c41a; font-size: 13px; }}
        #radarChart {{ width: 400px; height: 300px; margin: 16px auto; display: block; }}
        .footer {{ text-align: center; color: #aaa; margin-top: 40px; font-size: 13px; border-top: 1px solid #eee; padding-top: 16px; }}
        .ai-disclaimer {{ background: #fffbe6; border: 1px solid #ffe58f; border-radius: 6px; padding: 8px 14px; font-size: 12px; color: #888; margin-top: 8px; text-align: center; }}
    </style>
</head>
<body>
    <h1>职业规划报告</h1>
    <p style="text-align: center; color: #666;">生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
    <p class="ai-disclaimer">以下内容由 AI 生成，仅供参考，关键决策建议结合专业顾问意见。</p>

    <h2>基本信息</h2>
    <table>
        <tr><th>姓名</th><td>{report_data.get("student_name", "")}</td></tr>
        <tr><th>目标岗位</th><td>{report_data.get("job_name", "")}</td></tr>
        <tr><th>综合匹配度</th><td class="score">{report_data.get("overall_score", 0):.1f}%</td></tr>
    </table>

    <h2>匹配维度分析</h2>
    <div id="radarChart"></div>
    <div id="radarSvg" style="display:none;text-align:center">{radar_svg}</div>
    <table>
        <tr><th>维度</th><th>得分</th><th>说明</th></tr>
        <tr><td>基础要求</td><td>{dimensions.get("basic_requirements", {}).get("score", 0):.1f}%</td>
            <td>{dimensions.get("basic_requirements", {}).get("detail", "")}</td></tr>
        <tr><td>职业技能</td><td>{dimensions.get("professional_skills", {}).get("score", 0):.1f}%</td>
            <td>匹配{len(dimensions.get("professional_skills", {}).get("matched_skills", []))}项技能</td></tr>
        <tr><td>职业素养</td><td>{dimensions.get("professional_qualities", {}).get("score", 0):.1f}%</td>
            <td>{dimensions.get("professional_qualities", {}).get("detail", "")}</td></tr>
        <tr><td>发展潜力</td><td>{dimensions.get("development_potential", {}).get("score", 0):.1f}%</td>
            <td>{dimensions.get("development_potential", {}).get("detail", "")}</td></tr>
    </table>

    <h2>技能差距分析</h2>
    {skill_gaps_html}

    {chapters_html if chapters_html else f'''
    <h2>行动计划</h2>
    {action_plan_html}

    <h2>职业发展路径</h2>
    {career_path_html}
    '''}

    <p class="footer">— 报告由AI职业规划智能体生成 —</p>

    <script>
    var chart = echarts.init(document.getElementById('radarChart'));
    chart.setOption({{
        radar: {{
            indicator: {radar_indicator},
            shape: 'polygon',
            splitNumber: 5,
            radius: '65%',
        }},
        series: [{{
            type: 'radar',
            data: [{{
                value: {radar_values},
                name: '匹配得分',
                areaStyle: {{ color: 'rgba(24,144,255,0.2)' }},
                lineStyle: {{ color: '#1890ff', width: 2 }},
            }}],
        }}],
    }});
    </script>
</body>
</html>"""

        output_path = str(self.upload_dir / f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        # E-1: 导出质量自检
        content_len = len(html_content)
        if content_len < 2000:
            import logging as _log
            _log.getLogger(__name__).warning("[ExportQC] HTML内容过短（%d字符），报告可能不完整", content_len)

        return output_path

    def _build_radar_svg(self, dim_data: dict) -> str:
        """
        E-1: 服务端生成SVG雷达图，不依赖CDN/Canvas
        在导出Word时内嵌纯文本维度表格，PDF/HTML路径保留ECharts但附加SVG备用。
        """
        import math
        labels = list(dim_data.keys())
        values = list(dim_data.values())
        n = len(labels)
        if n == 0:
            return ""

        cx, cy, r = 200, 200, 140
        angles = [math.pi / 2 - 2 * math.pi * i / n for i in range(n)]

        # 网格线（5层）
        grid_lines = ""
        for layer in range(1, 6):
            pct = layer / 5
            pts = " ".join(
                f"{cx + r * pct * math.cos(a):.1f},{cy - r * pct * math.sin(a):.1f}"
                for a in angles
            )
            grid_lines += f'<polygon points="{pts}" fill="none" stroke="#dde0f0" stroke-width="1"/>\n'

        # 轴线
        axis_lines = ""
        for a in angles:
            x2, y2 = cx + r * math.cos(a), cy - r * math.sin(a)
            axis_lines += f'<line x1="{cx}" y1="{cy}" x2="{x2:.1f}" y2="{y2:.1f}" stroke="#dde0f0" stroke-width="1"/>\n'

        # 数据多边形
        data_pts = " ".join(
            f"{cx + r * (v / 100) * math.cos(a):.1f},{cy - r * (v / 100) * math.sin(a):.1f}"
            for v, a in zip(values, angles)
        )
        data_poly = f'<polygon points="{data_pts}" fill="rgba(24,144,255,0.25)" stroke="#1890ff" stroke-width="2"/>\n'

        # 标签
        label_els = ""
        for label, val, a in zip(labels, values, angles):
            lx = cx + (r + 20) * math.cos(a)
            ly = cy - (r + 20) * math.sin(a)
            anchor = "middle" if abs(math.cos(a)) < 0.3 else ("start" if math.cos(a) > 0 else "end")
            label_els += (
                f'<text x="{lx:.1f}" y="{ly:.1f}" text-anchor="{anchor}" '
                f'font-size="11" fill="#555">{label}({val:.0f})</text>\n'
            )

        svg = (
            f'<svg xmlns="http://www.w3.org/2000/svg" width="400" height="400" '
            f'viewBox="0 0 400 400">\n'
            f'{grid_lines}{axis_lines}{data_poly}{label_els}'
            f'</svg>'
        )
        return svg

    def _dim_table_text(self, dim_data: dict) -> str:
        """Word导出用：将维度分数转为文本表格行"""
        lines = ["维度评分一览："]
        for k, v in dim_data.items():
            bar = "█" * int(v / 10) + "░" * (10 - int(v / 10))
            lines.append(f"  {k:8s}  {bar}  {v:.0f}/100")
        return "\n".join(lines)


report_export_service = ReportExportService()
